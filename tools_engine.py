"""
Tools Engine — PDF utilities: merge, split, convert, protect, watermark, rotate.
All methods accept an optional progress callback cb(0-100).
"""
import os
import io
import shutil
import random
from pathlib import Path


def _check(name):
    try:
        __import__(name)
        return True
    except ImportError:
        return False


class ToolsEngine:
    def __init__(self):
        self.fitz_ok        = _check("fitz")
        self.pikepdf_ok     = _check("pikepdf")
        self.pillow_ok      = _check("PIL")
        self.pdf2docx_ok    = _check("pdf2docx")
        self.docx_ok        = _check("docx")
        self.pytesseract_ok = _check("pytesseract")   # Python package
        self.tesseract_ok   = self._check_tesseract()  # binary on PATH

    @staticmethod
    def _check_tesseract():
        import subprocess, sys

        # 1. Check common Windows install paths first (works without PATH restart)
        win_paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            r"C:\Tesseract-OCR\tesseract.exe",
            r"C:\tesseract\tesseract.exe",
        ]
        for wp in win_paths:
            if os.path.isfile(wp):
                # Also configure pytesseract to use this path
                try:
                    import pytesseract
                    pytesseract.pytesseract.tesseract_cmd = wp
                except ImportError:
                    pass
                return True

        # 2. Try PATH-based lookup
        for cmd in (["tesseract", "--version"], ["tesseract.exe", "--version"]):
            try:
                r = subprocess.run(cmd, capture_output=True, timeout=5)
                if r.returncode == 0:
                    return True
            except Exception:
                pass

        # 3. Ask pytesseract itself (it may know the path)
        try:
            import pytesseract
            ver = pytesseract.get_tesseract_version()
            if ver:
                return True
        except Exception:
            pass

        return False

    @staticmethod
    def _configure_pytesseract(pytesseract_module):
        """Point pytesseract at the Tesseract binary if it's in a common Windows path."""
        win_paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            r"C:\Tesseract-OCR\tesseract.exe",
            r"C:\tesseract\tesseract.exe",
        ]
        current = pytesseract_module.pytesseract.tesseract_cmd
        # Only override if current command isn't already a valid path
        if not os.path.isfile(str(current)):
            for wp in win_paths:
                if os.path.isfile(wp):
                    pytesseract_module.pytesseract.tesseract_cmd = wp
                    break

    def capabilities(self):
        return {
            "fitz": self.fitz_ok, "pikepdf": self.pikepdf_ok,
            "pillow": self.pillow_ok, "pdf2docx": self.pdf2docx_ok,
            "docx": self.docx_ok,
            "tesseract": self.tesseract_ok,
            "pytesseract": self.pytesseract_ok,
        }

    # ─── MERGE ────────────────────────────────────────────────────────────────
    def merge(self, paths, output, cb=None):
        cb = cb or (lambda p: None)
        cb(5)
        import fitz
        doc = fitz.open()
        for i, p in enumerate(paths):
            cb(5 + int(85 * i / max(len(paths), 1)))
            src = fitz.open(p)
            doc.insert_pdf(src)
            src.close()
        cb(92)
        doc.save(output, deflate=True)
        doc.close()
        cb(100)

    # ─── SPLIT ────────────────────────────────────────────────────────────────
    def split(self, path, output_dir, mode="each", ranges=None, cb=None):
        """
        mode='each'  → one PDF per page
        mode='range' → ranges like [(1,3),(4,6)]
        Returns list of output file paths.
        """
        cb = cb or (lambda p: None)
        cb(5)
        import fitz
        doc = fitz.open(path)
        n = len(doc)
        os.makedirs(output_dir, exist_ok=True)
        out_paths = []

        if mode == "each":
            for i in range(n):
                cb(5 + int(85 * i / max(n, 1)))
                d = fitz.open()
                d.insert_pdf(doc, from_page=i, to_page=i)
                op = os.path.join(output_dir, f"page_{i+1:03d}.pdf")
                d.save(op)
                d.close()
                out_paths.append(op)
        elif mode == "range" and ranges:
            for j, (s, e) in enumerate(ranges):
                cb(5 + int(85 * j / max(len(ranges), 1)))
                d = fitz.open()
                d.insert_pdf(doc, from_page=s - 1, to_page=e - 1)
                op = os.path.join(output_dir, f"pages_{s}_{e}.pdf")
                d.save(op)
                d.close()
                out_paths.append(op)

        doc.close()
        cb(100)
        return out_paths

    # ─── PDF → IMAGES ─────────────────────────────────────────────────────────
    def pdf_to_images(self, path, output_dir, dpi=150, fmt="PNG", cb=None):
        cb = cb or (lambda p: None)
        cb(5)
        import fitz
        doc = fitz.open(path)
        n = len(doc)
        os.makedirs(output_dir, exist_ok=True)
        out_paths = []
        for i, page in enumerate(doc):
            cb(5 + int(85 * i / max(n, 1)))
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat)
            ext = "jpg" if fmt.upper() in ("JPG", "JPEG") else fmt.lower()
            op = os.path.join(output_dir, f"page_{i+1:03d}.{ext}")
            if ext == "jpg":
                pix.save(op, jpg_quality=92)
            else:
                pix.save(op)
            out_paths.append(op)
        doc.close()
        cb(100)
        return out_paths

    # ─── IMAGES → PDF ─────────────────────────────────────────────────────────
    def images_to_pdf(self, paths, output, cb=None):
        cb = cb or (lambda p: None)
        cb(5)
        from PIL import Image
        imgs = []
        for i, p in enumerate(paths):
            cb(5 + int(75 * i / max(len(paths), 1)))
            img = Image.open(p)
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")
            imgs.append(img)
        if not imgs:
            raise RuntimeError("No valid images found.")
        cb(85)
        imgs[0].save(output, save_all=True, append_images=imgs[1:], resolution=150)
        cb(100)

    # ─── PDF → WORD ───────────────────────────────────────────────────────────
    def pdf_to_word(self, path, output, ocr=False, cb=None):
        cb = cb or (lambda p: None)
        cb(5)
        if not ocr:
            # ── Standard mode ────────────────────────────────────────────────
            # Step 1: try pdf2docx (preserves layout/tables better)
            pdf2docx_ok = False
            if self.pdf2docx_ok:
                try:
                    import logging
                    logging.getLogger("pdf2docx").setLevel(logging.ERROR)
                    from pdf2docx import Converter
                    cb(15)
                    cv = Converter(path)
                    cv.convert(output, start=0, end=None)
                    cv.close()
                    cb(60)
                    # Validate: open the docx and check it has real text
                    if self._docx_has_text(output):
                        pdf2docx_ok = True
                except Exception:
                    pass  # fall through to PyMuPDF extraction

            # Step 2: fallback — PyMuPDF direct text extraction
            if not pdf2docx_ok:
                cb(65)
                self._pymupdf_to_word(path, output, cb)
        else:
            self._ocr_to_word(path, output, cb)
        cb(100)

    @staticmethod
    def _docx_has_text(path):
        """Return True if the docx file contains at least some non-whitespace text."""
        try:
            from docx import Document
            doc = Document(path)
            text = " ".join(p.text for p in doc.paragraphs)
            return len(text.strip()) > 20
        except Exception:
            return False

    def _pymupdf_to_word(self, path, output, cb):
        """Rich text extraction: preserves headings, bold, italic, bullets via font analysis."""
        import fitz
        if not self.docx_ok:
            raise RuntimeError(
                "python-docx is not installed. Run install.bat or: pip install python-docx"
            )
        from docx import Document
        from docx.shared import Pt, Inches

        doc_pdf = fitz.open(path)
        n = len(doc_pdf)

        # Check if the PDF has any extractable text at all
        total_text = "".join(page.get_text() for page in doc_pdf).strip()
        if len(total_text) < 30:
            doc_pdf.close()
            raise RuntimeError(
                "This PDF contains no extractable text — it appears to be a "
                "scanned/image-based document. Please use OCR mode instead."
            )

        doc_out = Document()
        # Comfortable margins
        for section in doc_out.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1.25)
            section.right_margin  = Inches(1.25)

        doc_pdf = fitz.open(path)  # reopen after the check

        for page_i, page in enumerate(doc_pdf):
            cb(65 + int(25 * page_i / max(n, 1)))

            blocks = page.get_text(
                "dict",
                flags=fitz.TEXT_PRESERVE_LIGATURES | fitz.TEXT_PRESERVE_WHITESPACE
            ).get("blocks", [])

            # Collect all font sizes on the page to decide heading thresholds
            all_sizes = []
            for blk in blocks:
                if blk.get("type") != 0:
                    continue
                for line in blk.get("lines", []):
                    for span in line.get("spans", []):
                        sz = span.get("size", 11)
                        if sz > 0:
                            all_sizes.append(sz)
            body_size = sorted(all_sizes)[len(all_sizes) // 2] if all_sizes else 11  # median = body

            for blk in blocks:
                if blk.get("type") != 0:  # skip image blocks
                    continue

                for line in blk.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue

                    # Combine all text in this line to check its overall nature
                    line_text = "".join(sp.get("text", "") for sp in spans).strip()
                    if not line_text:
                        continue

                    # Dominant font size & flags for the line
                    dominant = max(spans, key=lambda s: s.get("size", 0))
                    dom_size  = dominant.get("size", 11)
                    dom_flags = dominant.get("flags", 0)
                    is_bold   = bool(dom_flags & (1 << 4))

                    # --- Classify as heading or body ---
                    if dom_size >= body_size * 1.6:          # very large → H1
                        doc_out.add_heading(line_text, level=1)
                    elif dom_size >= body_size * 1.25 or (dom_size >= body_size * 1.1 and is_bold):
                        doc_out.add_heading(line_text, level=2)
                    elif is_bold and dom_size >= body_size * 0.95 and len(line_text) < 80:
                        doc_out.add_heading(line_text, level=3)
                    else:
                        # Detect bullet characters at the start
                        import re
                        bullet_match = re.match(r'^([•·○◦►▸\-\*])\s+', line_text)
                        if bullet_match:
                            content = line_text[bullet_match.end():]
                            para = doc_out.add_paragraph(style="List Bullet")
                            _add_rich_run(para, spans, bullet_match.end())
                        else:
                            para = doc_out.add_paragraph()
                            _add_rich_run(para, spans)

            if page_i < n - 1:
                doc_out.add_page_break()

        doc_out.save(output)
        doc_pdf.close()

    def _ocr_to_word(self, path, output, cb):
        import fitz
        if not self.docx_ok:
            raise RuntimeError(
                "python-docx is not installed. "
                "Run install.bat or: pip install python-docx"
            )
        from docx import Document
        from docx.shared import Pt, Inches

        # Decide whether real OCR is available
        use_ocr = self.tesseract_ok and self.pytesseract_ok

        if use_ocr:
            try:
                import pytesseract
                from PIL import Image
                self._configure_pytesseract(pytesseract)
            except ImportError:
                use_ocr = False

        doc_out = Document()
        for section in doc_out.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1.25)
            section.right_margin  = Inches(1.25)

        doc_pdf = fitz.open(path)
        n = len(doc_pdf)

        if use_ocr:
            import pytesseract
            from PIL import Image
            self._configure_pytesseract(pytesseract)
            # 300 DPI for clean OCR (72 pt/inch → 300/72 ≈ 4.17x scale)
            scale = 300 / 72
            ocr_cfg = r"--oem 3 --psm 1"  # neural net + auto layout detection

            for i, page in enumerate(doc_pdf):
                cb(10 + int(75 * i / max(n, 1)))
                pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale))
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                try:
                    raw_text = pytesseract.image_to_string(img, config=ocr_cfg)
                except Exception as e:
                    raw_text = f"[OCR failed on page {i+1}: {e}]"

                _write_structured_ocr(doc_out, raw_text)

                if i < n - 1:
                    doc_out.add_page_break()
        else:
            # No OCR — check if PDF has digital text
            total_text = "".join(page.get_text() for page in doc_pdf).strip()
            if len(total_text) < 30:
                doc_pdf.close()
                if self.tesseract_ok and not self.pytesseract_ok:
                    raise RuntimeError(
                        "Tesseract is installed but the 'pytesseract' Python package is missing.\n"
                        "Run install.bat or: pip install pytesseract"
                    )
                raise RuntimeError(
                    "This PDF is scanned/image-based and contains no extractable text.\n"
                    "OCR (Tesseract) is required to convert it.\n"
                    "Install Tesseract from: https://github.com/UB-Mannheim/tesseract/wiki\n"
                    "Then run install.bat to install pytesseract."
                )
            # Fallback: use the rich digital extraction path
            doc_pdf.close()
            self._pymupdf_to_word(path, output, cb)
            return

        doc_out.save(output)
        doc_pdf.close()

    # ─── EDIT PDF — ANALYZE ───────────────────────────────────────────────────
    def edit_analyze(self, path):
        """Return text blocks per page for the in-app PDF editor."""
        import fitz
        doc = fitz.open(path)
        pages = []
        for i, page in enumerate(doc):
            blocks = []
            page_dict = page.get_text("dict")
            for block in page_dict.get("blocks", []):
                if block.get("type") != 0:
                    continue  # skip images
                block_text = ""
                fontsize = 11.0
                color = [0.0, 0.0, 0.0]
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        block_text += span.get("text", "")
                        fontsize = span.get("size", 11.0)
                        c = span.get("color", 0)
                        color = [
                            ((c >> 16) & 0xFF) / 255,
                            ((c >> 8) & 0xFF) / 255,
                            (c & 0xFF) / 255,
                        ]
                    block_text += "\n"
                block_text = block_text.strip()
                if not block_text:
                    continue
                bbox = block["bbox"]
                blocks.append({
                    "id": f"b{block['number']}",
                    "x0": round(bbox[0], 2), "y0": round(bbox[1], 2),
                    "x1": round(bbox[2], 2), "y1": round(bbox[3], 2),
                    "text": block_text,
                    "fontsize": round(fontsize, 1),
                    "color": color,
                    "page": i,
                })
            pages.append({
                "page": i,
                "width": round(page.rect.width, 2),
                "height": round(page.rect.height, 2),
                "blocks": blocks,
            })
        doc.close()
        return {"pages": pages, "total_pages": len(pages)}

    # ─── EDIT PDF — RENDER PAGE ───────────────────────────────────────────────
    def edit_render_page(self, path, page_num, dpi=120):
        """Render a single PDF page and return as base64 PNG string."""
        import fitz, base64
        doc = fitz.open(path)
        if page_num >= len(doc):
            raise RuntimeError(f"Page {page_num} out of range.")
        page = doc[page_num]
        mat = fitz.Matrix(dpi / 72, dpi / 72)
        pix = page.get_pixmap(matrix=mat)
        data = base64.b64encode(pix.tobytes("png")).decode()
        doc.close()
        return data

    # ─── EDIT PDF — APPLY EDITS ───────────────────────────────────────────────
    def edit_apply(self, path, output, edits, cb=None):
        """
        Apply text edits to a PDF.
        edits = list of {page, original, replacement, fontsize, color, bbox}
        """
        import fitz
        from collections import defaultdict
        cb = cb or (lambda p: None)
        cb(5)
        doc = fitz.open(path)
        n = len(doc)

        by_page = defaultdict(list)
        for edit in edits:
            by_page[int(edit["page"])].append(edit)

        for page_num in sorted(by_page.keys()):
            if page_num >= n:
                continue
            cb(5 + int(85 * page_num / max(n, 1)))
            page = doc[page_num]

            for edit in by_page[page_num]:
                original    = edit.get("original", "")
                replacement = edit.get("replacement", "")
                fontsize    = float(edit.get("fontsize", 11))
                raw_color   = edit.get("color", [0, 0, 0])
                color       = tuple(float(c) for c in raw_color)
                bbox        = edit.get("bbox")   # [x0,y0,x1,y1]

                if not original or original == replacement:
                    continue

                # ① Collect rects BEFORE redacting so we know where to re-insert
                insert_rect = None
                if bbox:
                    insert_rect = fitz.Rect(bbox[0], bbox[1], bbox[2], bbox[3])
                else:
                    # Search line by line to find the block area
                    for line in original.split("\n"):
                        line = line.strip()
                        if line:
                            rects = page.search_for(line)
                            if rects:
                                insert_rect = rects[0]
                                break

                # ② Redact original text — cover each line and the full bbox
                for line in original.split("\n"):
                    line = line.strip()
                    if line:
                        for r in page.search_for(line):
                            page.add_redact_annot(r, fill=(1, 1, 1))
                if bbox:
                    page.add_redact_annot(
                        fitz.Rect(bbox[0], bbox[1], bbox[2], bbox[3]),
                        fill=(1, 1, 1)
                    )
                page.apply_redactions()

                # ③ Insert replacement text at original block position
                if replacement.strip() and insert_rect:
                    point = fitz.Point(insert_rect.x0, insert_rect.y1 - 2)
                    page.insert_text(point, replacement,
                                     fontsize=fontsize, color=color)

        cb(90)
        doc.save(output, deflate=True)
        doc.close()
        cb(100)

    # ─── PROTECT ──────────────────────────────────────────────────────────────
    def protect(self, path, output, user_pw, owner_pw=None, cb=None):
        cb = cb or (lambda p: None)
        cb(20)
        import pikepdf
        with pikepdf.open(path) as pdf:
            pdf.save(
                output,
                encryption=pikepdf.Encryption(
                    user=user_pw,
                    owner=owner_pw or user_pw,
                    R=6,  # AES-256
                ),
            )
        cb(100)

    # ─── UNLOCK ───────────────────────────────────────────────────────────────
    def unlock(self, path, output, password="", cb=None):
        cb = cb or (lambda p: None)
        cb(20)
        import pikepdf
        try:
            with pikepdf.open(path, password=password) as pdf:
                pdf.save(output)
        except pikepdf.PasswordError:
            raise RuntimeError("Incorrect password — could not unlock PDF.")
        cb(100)

    # ─── SCAN LOOK (flatten every page to rasterised image) ───────────────────
    def flatten(self, path, output, style="scan", cb=None):
        """
        style='scan'  → grayscale, JPEG artefacts, slight rotation — looks scanned
        style='print' → colour, clean, non-editable (print-ready flat PDF)
        """
        cb = cb or (lambda p: None)
        cb(5)
        import fitz
        from PIL import Image, ImageEnhance, ImageFilter

        doc = fitz.open(path)
        n = len(doc)
        new_doc = fitz.open()

        for i, page in enumerate(doc):
            cb(5 + int(83 * i / max(n, 1)))
            pix = page.get_pixmap(matrix=fitz.Matrix(150 / 72, 150 / 72))
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            if style == "scan":
                img = img.convert("L")  # grayscale
                # Subtle scanner-drift rotation (±0.4°)
                img = img.rotate(
                    random.uniform(-0.4, 0.4), expand=False, fillcolor=255
                )
                # Slight softening (scanner glass effect)
                img = img.filter(ImageFilter.GaussianBlur(radius=0.35))
                # Contrast boost (like photocopier)
                img = ImageEnhance.Contrast(img).enhance(1.12)
                img = img.convert("RGB")

            buf = io.BytesIO()
            quality = 68 if style == "scan" else 90
            img.save(buf, "JPEG", quality=quality, optimize=True)
            buf.seek(0)

            # Re-wrap as single-page PDF
            tmp = fitz.open(stream=buf.read(), filetype="jpeg")
            new_doc.insert_pdf(fitz.open("pdf", tmp.convert_to_pdf()))

        new_doc.save(output, deflate=True)
        doc.close()
        cb(100)

    # ─── WATERMARK ────────────────────────────────────────────────────────────
    def watermark(self, path, output, text, opacity=0.15,
                  color=(180, 180, 180), angle=45, fontsize=30, cb=None):
        cb = cb or (lambda p: None)
        cb(5)
        import fitz
        doc = fitz.open(path)
        n = len(doc)
        r, g, b = [c / 255 for c in color]
        for i, page in enumerate(doc):
            cb(5 + int(85 * i / max(n, 1)))
            w, h = page.rect.width, page.rect.height
            for x in range(-int(w), int(w * 2), 220):
                for y in range(-int(h), int(h * 2), 130):
                    page.insert_text(
                        fitz.Point(x, y),
                        text,
                        fontsize=fontsize,
                        color=(r, g, b),
                        rotate=angle,
                        overlay=True,
                    )
        doc.save(output, deflate=True)
        doc.close()
        cb(100)

    # ─── ROTATE ───────────────────────────────────────────────────────────────
    def rotate(self, path, output, angle=90, pages="all", cb=None):
        cb = cb or (lambda p: None)
        cb(10)
        import pikepdf
        with pikepdf.open(path) as pdf:
            n = len(pdf.pages)
            if pages == "all":
                idxs = range(n)
            else:
                idxs = [
                    int(p.strip()) - 1
                    for p in str(pages).split(",")
                    if p.strip().isdigit()
                ]
            for i in idxs:
                if 0 <= i < n:
                    cb(10 + int(80 * (i / max(n, 1))))
                    page = pdf.pages[i]
                    cur = int(page.get("/Rotate", 0))
                    page["/Rotate"] = (cur + angle) % 360
            pdf.save(output)
        cb(100)


# ── Module-level helpers for PDF→Word formatting ─────────────────────────────

def _add_rich_run(para, spans, skip_chars=0):
    """Add formatted runs to a python-docx paragraph from PyMuPDF span list."""
    from docx.shared import Pt
    combined = "".join(sp.get("text", "") for sp in spans)
    # If we need to strip a bullet prefix, work on combined text
    if skip_chars:
        combined = combined[skip_chars:]
        run = para.add_run(combined)
        run.font.size = Pt(11)
        return

    for sp in spans:
        text = sp.get("text", "")
        if not text:
            continue
        run = para.add_run(text)
        flags    = sp.get("flags", 0)
        run.bold   = bool(flags & (1 << 4))
        run.italic = bool(flags & (1 << 1))
        size = sp.get("size", 11)
        run.font.size = Pt(max(7, min(size, 28)))


def _write_structured_ocr(doc_out, text):
    """Parse raw OCR text and write structured paragraphs into a docx Document."""
    import re
    from docx.shared import Pt

    # Pre-clean: drop lines that are clearly OCR noise
    # (< 4 printable chars, or > 40 % non-alphanumeric / non-space chars)
    def _is_noise(line):
        s = line.strip()
        if len(s) < 4:
            return True
        alnum_space = sum(1 for c in s if c.isalnum() or c.isspace())
        return alnum_space / max(len(s), 1) < 0.55

    lines = [l for l in text.splitlines() if not _is_noise(l)]

    prev_ended_colon = False   # track whether previous line ended with ":"
    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        i += 1

        if not stripped:
            prev_ended_colon = False
            continue

        # ── Bullet detection ──────────────────────────────────────────────────
        bullet_m = re.match(r'^([•·○◦►▸•‣◦\[\]]|-|\*)\s+(.+)$', stripped)
        if bullet_m:
            content = bullet_m.group(2)
            p = doc_out.add_paragraph(content, style="List Bullet")
            prev_ended_colon = False
            continue

        # ── Heading detection heuristics ──────────────────────────────────────
        has_mid_colon = ":" in stripped and not stripped.endswith(":")
        is_short      = len(stripped) <= 70
        no_period     = not stripped.endswith((".", ",", ";"))
        all_caps      = stripped.isupper() and len(stripped) > 3
        title_ish     = stripped.istitle() and not has_mid_colon

        # Suppress heading if previous line ended with ":" — these are list items
        if is_short and no_period and not has_mid_colon and (all_caps or title_ish) and not prev_ended_colon:
            level = 1 if all_caps else 2
            doc_out.add_heading(stripped, level=level)
            prev_ended_colon = False
            continue

        # If it looked like a potential heading but was suppressed by colon-context,
        # write it as a List Bullet instead
        if is_short and no_period and not has_mid_colon and (all_caps or title_ish) and prev_ended_colon:
            p = doc_out.add_paragraph(stripped, style="List Bullet")
            prev_ended_colon = False
            continue

        # ── Bold sub-heading: short, ends with colon ──────────────────────────
        if is_short and stripped.endswith(":") and len(stripped) < 60:
            p = doc_out.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            prev_ended_colon = True
            continue

        # ── Normal paragraph ──────────────────────────────────────────────────
        p = doc_out.add_paragraph(stripped)
        p.style.font.size = Pt(11)
        prev_ended_colon = stripped.endswith(":")
